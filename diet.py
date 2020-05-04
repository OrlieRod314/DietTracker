''' File: diet.py
    Author: Orlando Rodriguez
    Purpose: Runs the diet tracker
'''
import os, sys
import csv
import argparse

from entry import make_entries, Entry
from diet_lims import Lim
from colorama import Fore, Style, init
from docx import *

module_name = "DietTracker: Track nutrition and make better choices"
__version__ = "1.0"    

def write_report(entries, limit, macro_sums, args):
    outputs = 'outputs/'
    output_file_name = args.dataset_file_name + '_output.txt'
    sum_difs = [0, 0, 0, 0]
    
    f = open(outputs + output_file_name, 'w')
    f.write("Nutrition report for: " + args.dataset_file_name)
    f.write("\nSurplus/Deficiency table:\nCals Carbs Protein Fat\n")
    for i in range(len(limit.get_lims())):
        sum_difs[i] = limit.get_lims()[i] - macro_sums[i]
    f.write(str(sum_difs))
    f.close()

def write_report_docx(entries, limit, macro_sums, args):
    sum_difs = [0, 0, 0, 0]
    
    document = Document()
    document.add_heading('Nutrition Report')
    p = document.add_paragraph('Here are the nutritional surpluses and deficiencies for a ')
    p.add_run(str(limit.get_profile()[0])).bold = True
    p.add_run('-year old ')
    p.add_run(limit.get_profile()[1]).bold = True
    p.add_run(' following the diet detailed in the ')
    p.add_run(args.dataset_file_name).bold = True
    p.add_run(' dataset.')
    
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Cals'
    hdr_cells[1].text = 'Carbs'
    hdr_cells[2].text = 'Prot'
    hdr_cells[3].text = 'Fat'
    row_cells = table.add_row().cells
    for i in range(len(limit.get_lims())):
        row_cells[i].text = str(limit.get_lims()[i] - macro_sums[i])
    outputs = 'outputs/'
    output_file_name = args.dataset_file_name + '_output.docx'
    document.add_page_break()
    document.save(outputs + output_file_name)

def main():
    parser = argparse.ArgumentParser(description = f"{module_name} (Version {__version__})")
    parser.add_argument('age', 
                            metavar = 'AGE', type=int,
                            help='User age')
    parser.add_argument('sex', 
                            metavar = 'SEX', type=str,
                            help='User sex')
    parser.add_argument('dataset_file_name', 
                            metavar = 'DATASET_FILE_NAME', type=str,
                            help='Input file that information will be read from')
    parser.add_argument('-kg', '--ketogenic', 
                            action='store_true', dest='kg', default = False, 
                            help='Offers recommendation based on ketogenic diet')
    parser.add_argument('-docx', '--word',
                            action='store_true', default = False,
                            help='Writes report to .docx file instead of .txt')
    parser.add_argument('-ver', '--version',
                            action='version',  version=__version__,
                            help='Display version information and dependencies.')
    parser.add_argument('-nocol', '--nocolor',
                            action='store_true', default = False, 
                            help='Disables color in terminal')
    detail = parser.add_mutually_exclusive_group()
    detail.add_argument('-q', '--quiet',
                            action='store_true', 
                            help='Print quiet')
    detail.add_argument('-v', '--verbose',
                            action='store_true', 
                            help='Print verbose')
    args = parser.parse_args()
    
    lines = []
    macro_sums = [0, 0, 0, 0]
    datasets = 'datasets/'
    dataset_file_name = args.dataset_file_name + '.csv'
    
    if not args.nocolor:
        init()
    if args.verbose:
        if args.nocolor:
            print("Reading from file " + dataset_file_name)
        else:
            print(Style.BRIGHT + Fore.GREEN + "Reading from file " + dataset_file_name)
            print(Style.RESET_ALL)
    
    with open(datasets + dataset_file_name, 'r') as csv_file:
        csv_reader = csv.reader(csv_file)
        for line in csv_reader:
            for i in range(1, 5):
                macro_sums[i - 1] += int(line[i])
            lines.append(line)
    
    entries = make_entries(lines)
    limit = Lim(args.age, args.sex)
    
    if not args.quiet:
        print("Entries:\nName\tCals\tCarbs\tProtein\t\tFat\n")
        for entry in entries:
            print(entry)
        if args.verbose:
            if args.nocolor:
                print("Macro sums: " + str(macro_sums))
                print("Limit: " + str(limit.get_lims()))
            else:
                print(Style.BRIGHT + Fore.CYAN + "Macro sums: " + str(macro_sums))
                print(Style.BRIGHT + Fore.CYAN + "Limit: " + str(limit.get_lims()))
                print(Style.RESET_ALL)
    
    if args.word:
        write_report_docx(entries, limit, macro_sums, args)
    else:
        write_report(entries, limit, macro_sums, args)

    if not args.quiet:
        if args.nocolor:
            print("Output Created")
        else:    
            print(Style.BRIGHT + Fore.GREEN + "Output Created")
            print(Style.RESET_ALL)
    
if __name__ == "__main__":
    main()